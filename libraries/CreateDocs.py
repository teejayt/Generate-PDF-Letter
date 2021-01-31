# +
from docx import Document
from docx.shared import Inches
from docx2pdf import convert


class CreateDocs:
    def createWordDocument(self, filePath):
        #Create a new Mortgage Payoff Letter -https://www.pinterest.com/pin/609252655826930727/
        document = Document()
        p = document.add_paragraph('Attention Payoff Department')
        p = document.add_paragraph('')
        p = document.add_paragraph('')
        p.add_run('Re:\tPayoff Request---C123490\n').bold = True
        p.add_run('\tMortgagor: Joe Bloggs\n\n').bold = True
        p.add_run('\tMortgage No.: PQ2990003\n').bold = True
        p.add_run('\tPremises: 145 Joe Smith Way, California\n\n').bold = True
        p.add_run('\tAnticipitated Closing Date: 24/01/2021').bold = True

        document.add_paragraph('Dear Sir or Madam:')
        document.add_paragraph('Please be advised that our clients wish to satisy the above mortgage. Kindly forward us a Satisfaction letter.')
        document.add_paragraph('If your mortgage is a participation/equity source loan, kindly confirm that the assets of said loan are frozen upon your receipt of this letter.')
        document.add_paragraph('Very truly yours,')
        document.add_paragraph('---------------------------------------------')
        document.save(filePath)
        
    def createPDF(self, wordDoc, pdfDoc):
        #will work for both Windows and MacOS
        convert(wordDoc, pdfDoc)
